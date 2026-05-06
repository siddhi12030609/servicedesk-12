from docx import Document
from docx.shared import Inches
import requests
import os

# -------- FILE PATHS (USE FULL PATH - VERY IMPORTANT) --------
doc_path = r"C:\Users\Sidddhi\Desktop\ServiceDeskBot_Project_Guide.docx"
final_doc_path = r"C:\Users\Sidddhi\Desktop\ServiceDeskBot_Project_Full.docx"
img_path = r"C:\Users\Sidddhi\Desktop\ServiceDeskBot_Architecture.png"

# -------- LOAD EXISTING DOCUMENT --------
doc = Document(doc_path)

# -------- ADD ARCHITECTURE SECTION --------
doc.add_page_break()
doc.add_heading('Technical Architecture Diagram', level=1)
doc.add_paragraph(
    "Below is the technical architecture diagram of the Service Desk Chatbot system, showing "
    "the flow from user frontend to Auth0 authentication, FastAPI backend, role-based approvals, "
    "notifications, and database logging."
)

# -------- DOWNLOAD IMAGE SAFELY --------
diagram_url = "https://i.imgur.com/3UKq6xK.png"

try:
    response = requests.get(diagram_url, timeout=10)
    response.raise_for_status()  # Raises error if bad response

    # Check if response is actually an image
    if "image" in response.headers.get("Content-Type", ""):
        with open(img_path, "wb") as f:
            f.write(response.content)

        # Ensure file is not empty
        if os.path.getsize(img_path) > 0:
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph("Figure 1: Technical Architecture of Service Desk Chatbot")
            print("✅ Image inserted successfully.")
        else:
            print("❌ Downloaded file is empty.")
    else:
        print("❌ The URL did not return a valid image.")

except Exception as e:
    print("❌ Error downloading image:", e)

# -------- SAVE FINAL DOCUMENT --------
doc.save(final_doc_path)
print("✅ Final Word document created at Desktop.")